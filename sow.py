import streamlit as st
from datetime import date
import io
import re
import os
import time 
import requests
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- FILE PATHING & DIAGRAM MAPPING ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(BASE_DIR, "diagrams")

AWS_PN_LOGO = os.path.join(ASSETS_DIR, "aws partner logo.jpg")
ONETURE_LOGO = os.path.join(ASSETS_DIR, "oneture logo1.jpg")
AWS_ADV_LOGO = os.path.join(ASSETS_DIR, "aws advanced logo.jpg")

SOW_COST_TABLE_MAP = { 
    "L1 Support Bot POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Beauty Advisor POC SOW": { 
        "poc_cost": "4,525.66 USD + 200 USD (Amazon Bedrock Cost) = 4,725.66", 
        "prod_cost": "4,525.66 USD + 1,175.82 USD (Amazon Bedrock Cost) = 5,701.48" 
    }, 
    "Ready Search POC Scope of Work Document":{ "poc_cost": "2,641.40 USD" }, 
    "AI based Image Enhancement POC SOW": { "poc_cost": "2,814.34 USD" }, 
    "AI based Image Inspection POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Gen AI for SOP POC SOW": { "poc_cost": "2,110.30 USD" }, 
    "Project Scope Document": { "prod_cost": "2,993.60 USD" }, 
    "Gen AI Speech To Speech": { "prod_cost": "2,124.23 USD" }, 
    "PoC Scope Document": { "amazon_bedrock": "1,000 USD", "total": "$ 3,150" }
}

CALCULATOR_LINKS = {
    "L1 Support Bot POC SOW": "https://calculator.aws/#/estimate?id=211ea64cba5a8f5dc09805f4ad1a1e598ef5238b",
    "Ready Search POC Scope of Work Document": "https://calculator.aws/#/estimate?id=f8bc48f1ae566b8ea1241994328978e7e86d3490",
    "AI based Image Enhancement POC SOW": "https://calculator.aws/#/estimate?id=9a3e593b92b796acecf31a78aec17d7eb957d1e5",
    "Beauty Advisor POC SOW": "https://calculator.aws/#/estimate?id=3f89756a35f7bac7b2cd88d95f3e9aba9be9b0eb",
    "AI based Image Inspection POC SOW": "https://calculator.aws/#/estimate?id=72c56f93b0c0e101d67a46af4f4fe9886eb93342",
    "Gen AI for SOP POC SOW": "https://calculator.aws/#/estimate?id=c21e9b242964724bf83556cfeee821473bb935d1",
    "Project Scope Document": "https://calculator.aws/#/estimate?id=37339d6e34c73596559fe09ca16a0ac2ec4c4252",
    "Gen AI Speech To Speech": "https://calculator.aws/#/estimate?id=8444ae26e6d61e5a43e8e743578caa17fd7f3e69",
    "PoC Scope Document": "https://calculator.aws/#/estimate?id=420ed9df095e7824a144cb6c0e9db9e7ec3c4153"
}

SOW_DIAGRAM_MAP = {
    "L1 Support Bot POC SOW": os.path.join(ASSETS_DIR, "L1 Support Bot POC SOW.png"),
    "Beauty Advisor POC SOW": os.path.join(ASSETS_DIR, "Beauty Advisor POC SOW.png"),
    "Ready Search POC Scope of Work Document": os.path.join(ASSETS_DIR, "Ready Search POC Scope of Work Document.png"),
    "AI based Image Enhancement POC SOW": os.path.join(ASSETS_DIR, "AI based Image Enhancement POC SOW.png"),
    "AI based Image Inspection POC SOW": os.path.join(ASSETS_DIR, "AI based Image Inspection POC SOW.png"),
    "Gen AI for SOP POC SOW": os.path.join(ASSETS_DIR, "Gen AI for SOP POC SOW.png"),
    "Project Scope Document": os.path.join(ASSETS_DIR, "Project Scope Document.png"),
    "Gen AI Speech To Speech": os.path.join(ASSETS_DIR, "Gen AI Speech To Speech.png"),
    "PoC Scope Document": os.path.join(ASSETS_DIR, "PoC Scope Document.png")
}

st.set_page_config(page_title="GenAI SOW Architect", layout="wide", page_icon="📄")

st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .stakeholder-header { background-color: #f1f5f9; padding: 8px 12px; border-radius: 6px; margin-top: 10px; font-weight: bold; border-left: 4px solid #3b82f6; }
    .sow-preview { background-color: white; padding: 40px; border-radius: 12px; border: 1px solid #e2e8f0; line-height: 1.7; color: #000000; font-family: "Times New Roman", Times, serif; }
    .sow-preview a { color: #3b82f6; text-decoration: underline; }
    .timeline-table { border-collapse: collapse; width: 100%; margin-top: 10px; border: 1px solid black; }
    .timeline-table th, .timeline-table td { border: 1px solid black; padding: 8px; text-align: center; font-size: 10pt; color: #000000; }
    .shaded-cell { background-color: #D9D9D9 !important; color: transparent !important; }
    </style>
    """, unsafe_allow_html=True)

# --- DOCX UTILS ---
def set_cell_shading(cell, color="D9D9D9"):
    """
    Sets the background shading of a table cell.
    Explicitly sets w:val="clear" to ensure the fill color renders correctly.
    """
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000EE') 
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(c); rPr.append(u); new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t); hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_docx_logic(text_content, branding, sow_name, timeline_df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Page 1: Cover
    p_logo = doc.add_paragraph()
    if os.path.exists(AWS_PN_LOGO): p_logo.add_run().add_picture(AWS_PN_LOGO, width=Inches(1.6))
    doc.add_paragraph("\n" * 3)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(branding['sow_name']); run.font.size = Pt(26); run.bold = True
    stitle = doc.add_paragraph(); stitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stitle.add_run("Scope of Work Document").font.size = Pt(14)
    doc.add_paragraph("\n" * 4)
    
    l_table = doc.add_table(rows=1, cols=3); l_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if branding.get("customer_logo_bytes"):
        l_table.rows[0].cells[0].paragraphs[0].add_run().add_picture(io.BytesIO(branding["customer_logo_bytes"]), width=Inches(1.8))
    if os.path.exists(ONETURE_LOGO):
        l_table.rows[0].cells[1].paragraphs[0].add_run().add_picture(ONETURE_LOGO, width=Inches(2.2))
    if os.path.exists(AWS_ADV_LOGO):
        l_table.rows[0].cells[2].paragraphs[0].add_run().add_picture(AWS_ADV_LOGO, width=Inches(1.8))
    
    doc.add_paragraph("\n" * 3)
    dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt.add_run(branding["doc_date_str"]).bold = True
    doc.add_page_break()

    # Page 2: TOC
    h_toc = doc.add_heading("Table of Contents", level=1)
    for run in h_toc.runs: run.bold = True; run.font.color.rgb = RGBColor(0,0,0)
    toc_items = ["2 Project Overview", "  2.1 Objective", "  2.2 Project Sponsor(s) / Stakeholder(s) / Project Team", "  2.3 Assumptions & Dependencies", "  2.4 PoC Success Criteria", "3 Scope of Work - Technical Project Plan", "4 Solution Architecture / Architectural Diagram", "5 RESOURCES & COST ESTIMATES"]
    for item in toc_items: doc.add_paragraph(item)
    doc.add_page_break()

    # Section Headers
    headers_map = {"2": "Project Overview", "3": "Scope of Work - Technical Project Plan", "4": "Solution Architecture / Architectural Diagram", "5": "RESOURCES & COST ESTIMATES"}
    lines = text_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line: i += 1; continue
        clean_line = re.sub(r'#+\s*', '', line).strip()
        # Remove asterisks for cleaner processing
        clean_line_no_ast = clean_line.replace('*', '')
        upper = clean_line_no_ast.upper()
        
        current_id = None
        for h_id, h_title in headers_map.items():
            if re.match(rf"^{h_id}[\.\s]+.*{re.escape(h_title.split()[0].upper())}", upper):
                current_id = h_id; break

        if current_id:
            if current_id != "2": doc.add_page_break()
            h = doc.add_heading(clean_line_no_ast.upper(), level=1)
            for run in h.runs: run.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)
            
            if current_id == "4":
                diag = SOW_DIAGRAM_MAP.get(sow_name)
                if diag and os.path.exists(diag):
                    doc.add_picture(diag, width=Inches(5.5))
                    p_cap = doc.add_paragraph(f"{sow_name} – Architecture Diagram")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue
            
        if line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            # Skip Markdown timeline tables as we generate them manually
            # Also checking slightly wider window to avoid hallucinations
            if any("Development Timelines" in lines[max(0, i-j)] for j in range(1, 4)):
                i += 1
                continue

            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'): table_lines.append(lines[i]); i += 1
            if len(table_lines) >= 3:
                cols = [c.strip().replace('*','') for c in table_lines[0].split('|') if c.strip()]
                t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"
                for idx, h_text in enumerate(cols):
                    t.rows[0].cells[idx].paragraphs[0].add_run(h_text).bold = True
                for row_line in table_lines[2:]:
                    cells_data = [c.strip().replace('*','') for c in row_line.split('|') if c.strip()]
                    r = t.add_row().cells
                    for idx, c_text in enumerate(cells_data): 
                        if idx < len(r): 
                            if "link" in c_text.lower():
                                add_hyperlink(r[idx].paragraphs[0], "Link", CALCULATOR_LINKS.get(sow_name, "https://calculator.aws/"))
                            else:
                                r[idx].paragraphs[0].add_run(c_text)
            continue

        # Enhanced detection for Timeline section
        is_timeline_header = "Development Timelines" in clean_line_no_ast or "Development Timeline" in clean_line_no_ast
        is_timeline_placeholder = "[TIMELINE_PLACEHOLDER]" in line

        if is_timeline_header or is_timeline_placeholder:
            # If we matched the header text, add it to the doc
            if is_timeline_header:
                p_tl = doc.add_paragraph()
                p_tl.add_run("Development Timelines:").bold = True
            
            # --- INSERT CORRECT DYNAMIC TABLE ---
            cols = timeline_df.columns.tolist()
            t = doc.add_table(rows=1, cols=len(cols)); t.style = "Table Grid"
            for idx, h_text in enumerate(cols): t.rows[0].cells[idx].paragraphs[0].add_run(h_text).bold = True
            for _, row in timeline_df.iterrows():
                r = t.add_row().cells
                for idx, col_name in enumerate(cols):
                    val = str(row[col_name])
                    if val.strip().upper() == "X": 
                        set_cell_shading(r[idx])
                    else: 
                        r[idx].paragraphs[0].add_run(val)
            
            # --- CONSUME/SKIP HALLUCINATED TABLE/PLACEHOLDER ---
            # Peek ahead and eat lines that look like a table or placeholder text
            while i + 1 < len(lines):
                next_line_check = lines[i+1].strip()
                # Stop skipping if we hit a new header
                if next_line_check.startswith('#') or "Solution Architecture" in next_line_check:
                    break
                # Skip tables, empty lines, or placeholder text
                if not next_line_check or next_line_check.startswith('|') or "inserted manually" in next_line_check or "Description" in next_line_check or "placeholder" in next_line_check.lower() or "[TIMELINE_PLACEHOLDER]" in next_line_check:
                    i += 1
                else:
                    break
            
            i += 1; continue

        # Identify sub-headings like 2.1, A., B.
        is_sub = re.match(r'^[A-Z]\.\s+', clean_line_no_ast) or re.match(r'^\d+\.\d+\s+', clean_line_no_ast) or re.match(r'^\d+\.\s+', clean_line_no_ast)
        
        if line.startswith('## ') or line.startswith('### ') or is_sub: 
            lvl = 2 
            h = doc.add_heading(clean_line_no_ast, level=lvl)
            for run in h.runs: run.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.strip().startswith('o '):
            p_b = doc.add_paragraph(style="List Bullet")
            p_b.add_run(line.strip()[2:].strip().replace('*', ''))
        elif line.strip().startswith('§ '):
            p_b = doc.add_paragraph(style="List Bullet")
            p_b.paragraph_format.left_indent = Inches(0.5)
            p_b.add_run(line.strip()[2:].strip().replace('*', ''))
        elif line.startswith('- ') or line.startswith('* '):
            p_b = doc.add_paragraph(style="List Bullet")
            p_b.add_run(re.sub(r'^[\-\]\s]+', '', line).strip().replace('*', ''))
        else:
            p_n = doc.add_paragraph()
            p_n.add_run(clean_line_no_ast)
        i += 1
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def call_gemini_with_retry(payload, api_key_input=""):
    apiKey = api_key_input if api_key_input else ""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={apiKey}"
    delays = [1, 2, 4, 8, 16]
    for attempt in range(len(delays)):
        try:
            res = requests.post(url, json=payload, timeout=30)
            if res.status_code == 200: return res, None
            if res.status_code in [503, 429]: time.sleep(delays[attempt]); continue
            return None, f"API Error {res.status_code}: {res.text}"
        except Exception: time.sleep(delays[attempt])
    return None, "Model overloaded."

# --- INITIALIZATION ---
def init_state():
    if 'generated_sow' not in st.session_state: st.session_state.generated_sow = ""
    if 'stakeholders' not in st.session_state:
        st.session_state.stakeholders = {
            "Partner": pd.DataFrame([{"Name": "Gaurav Kankaria", "Title": "Head of Analytics & ML", "Email": "gaurav.kankaria@oneture.com"}]),
            "Customer": pd.DataFrame([{"Name": "Prafful Mathur", "Title": "Lead Engineer", "Email": "prafful.mathur@nykaa.com"}]),
            "AWS": pd.DataFrame([{"Name": "Anubhav Sood", "Title": "AWS Account Executive", "Email": "anbhsood@amazon.com"}]),
            "Escalation": pd.DataFrame([{"Name": "Omkar Dhavalikar", "Title": "AI/ML Lead", "Email": "omkar.dhavalikar@oneture.com"}, {"Name": "Gaurav Kankaria", "Title": "Head of Analytics and AIML", "Email": "gaurav.kankaria@oneture.com"}])
        }
    if 'timeline_phases' not in st.session_state:
        st.session_state.timeline_phases = pd.DataFrame([
            {"Phase": "Infrastructure Setup", "Task": "Setup AWS Services & gather documents", "Wk1": "X", "Wk2": ""},
            {"Phase": "Core Workflow", "Task": "SOP Processing, logic & integration", "Wk1": "", "Wk2": "X"},
            {"Phase": "UI & Testing", "Task": "Streamlit build & Feedback loops", "Wk1": "", "Wk2": ""},
            {"Phase": "Final Review", "Task": "Documentation and Sign-off", "Wk1": "", "Wk2": ""}
        ])

init_state()

# --- SIDEBAR ---
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/artificial-intelligence.png", width=60)
    st.title("Architect Pro")
    api_key = st.text_input("Gemini API Key", type="password")
    st.divider()
    st.header("📋 1. Project Intake")
    sow_opts = ["1. L1 Support Bot POC SOW", "2. Beauty Advisor POC SOW", "3. Ready Search POC Scope of Work Document", "4. AI based Image Enhancement POC SOW", "5. AI based Image Inspection POC SOW", "6. Gen AI for SOP POC SOW", "7. Project Scope Document", "8. Gen AI Speech To Speech", "9. PoC Scope Document", "Other (Custom Use Case)"]
    solution_type = st.selectbox("1.1 Solution Type", sow_opts)
    if solution_type == "Other (Custom Use Case)":
        custom_name = st.text_input("Enter Custom Use Case Name:", placeholder = "Enter Custom Use Case")
        sow_key = custom_name
    else:
        sow_key = solution_type.split(". ", 1)[1] if ". " in solution_type else solution_type
        

    engagement_type = st.selectbox("1.2 Engagement Type", ["Proof of Concept (PoC)", "Pilot", "MVP", "Production Rollout", "Assessment / Discovery", "Support"])
    industry_type = st.selectbox("1.3 Industry / Domain", ["Retail / E-commerce", "BFSI", "Manufacturing", "Telecom", "Healthcare", "Energy / Utilities", "Logistics", "Media", "Government", "Other (specify)"])
    final_industry = st.text_input("Specify Industry:") if industry_type == "Other (specify)" else industry_type
    if st.button("🗑️ Reset All", use_container_width=True): 
        for key in list(st.session_state.keys()): del st.session_state[key]
        init_state(); st.rerun()

# --- MAIN UI ---
st.title("🚀 GenAI Scope of Work Architect")
st.header("📸 Cover Page Branding")
col_cov1, col_cov2 = st.columns(2)
with col_cov1: customer_logo = st.file_uploader("Upload Customer Logo", type=["png", "jpg", "jpeg"])
with col_cov2: doc_date = st.date_input("Document Date", date.today())
st.divider()

st.header("2. Project Overview Section")
biz_objective = st.text_area("2.1 Business Objective", placeholder="What business problem is the customer trying to solve?", height=100)
sel_outcomes = st.multiselect("2.2 Key Outcomes Expected", ["Reduce manual effort", "Improve accuracy / quality", "Faster turnaround time", "Cost reduction", "Revenue uplift", "Compliance improvement", "Better customer experience", "Scalability validation", "Other (specify)"], default=["Improve accuracy / quality", "Cost reduction"])
st.subheader("👥 2.3 Stakeholders Information")
st.session_state.stakeholders["Partner"] = st.data_editor(st.session_state.stakeholders["Partner"], num_rows="dynamic", key="ed_p")
st.session_state.stakeholders["Customer"] = st.data_editor(st.session_state.stakeholders["Customer"], num_rows="dynamic", key="ed_c")
st.session_state.stakeholders["AWS"] = st.data_editor(st.session_state.stakeholders["AWS"], num_rows="dynamic", key="ed_a")
st.session_state.stakeholders["Escalation"] = st.data_editor(st.session_state.stakeholders["Escalation"], num_rows="dynamic", key="ed_e")
st.divider()

st.header("📋 3. Assumptions & Dependencies")
dep_opts = ["Sample data availability", "Historical data availability", "Design / business guidelines finalized", "API access provided", "User access to AWS account", "SME availability for validation", "Network / VPC access", "Security approvals"]
sel_deps = [opt for opt in dep_opts if st.checkbox(opt, key=f"dep_{opt}")]
st.subheader("📊 3.2 Data Characteristics")
data_types = st.multiselect("Data involved:", ["Images", "Text", "PDFs / Documents", "Audio", "Video", "Structured tables", "APIs / Streams"])
data_meta = {}
for dt in data_types:
    with st.expander(f"⚙️ {dt} Details", expanded=True):
        c1, c2, c3 = st.columns(3)
        data_meta[dt] = {"Size": c1.text_input(f"{dt} Avg Size", "2 MB"), "Format": c2.text_input(f"{dt} Formats", "PDF"), "Vol": c3.text_input(f"{dt} Volume", "100/day")}
sel_ass = [opt for opt in ["PoC only, not production-grade", "Limited data volume", "Rule-based logic acceptable initially", "Manual review for edge cases", "No real-time SLA commitments"] if st.checkbox(opt, key=f"ass_{opt}")]
custom_ass = st.text_input("Other Assumptions:", key="custom_ass_in")
st.divider()

st.header("🎯 4. PoC Success Criteria")
sel_dims = st.multiselect("Dimensions:", ["Accuracy", "Latency", "Usability", "Explainability", "Coverage", "Cost efficiency", "Integration readiness"], default=["Accuracy", "Cost efficiency"])
val_req = st.radio("Validation Strategy:", ["Yes – customer validation required", "No – internal validation sufficient"])
st.divider()

st.header("🛠️ 5. Scope of Work")
sel_caps = [c for c in ["Upload / Ingestion", "Processing / Inference", "Metadata extraction", "Scoring / Recommendation", "Feedback loop", "UI display"] if st.checkbox(c, value=True, key=f"cap_{c}")]
custom_cap = st.text_input("Add Custom Step:", key="custom_cap_in")
sel_ints = st.multiselect("Integrations:", ["Internal databases", "External APIs", "CRM", "ERP", "Search engine", "Data warehouse", "None"], default=["None"])
st.divider()

st.header("🏢 6. Architecture & AWS Services")
compute_choices = st.multiselect("Compute Options:", ["AWS Lambda", "Step Functions", "Amazon ECS / EKS(future)", "Hybrid"], default=["AWS Lambda", "Step Functions"])
ai_svcs = st.multiselect("AI Services:", ["Amazon Bedrock", "Amazon SageMaker", "Rekognition", "Textract", "Comprehend", "Transcribe", "Translate"], default=["Amazon Bedrock"])
st_svcs = st.multiselect("Storage:", ["Amazon S3", "DynamoDB", "OpenSearch", "RDS", "Vector DB (OpenSearch / Aurora PG)"], default=["Amazon S3"])
ui_layer = st.selectbox("UI Layer:", ["Streamlit on S3", "CloudFront + Static UI", "Internal demo only", "No UI (API only)"], index=0)
st.divider()

st.header("⚙️ 7. Non-Functional Requirements")
perf = st.selectbox("Performance Profile:", ["Batch", "Near real-time", "Real-time"], index=1)
sec = st.multiselect("Security Controls:", ["IAM-based access", "Encryption at rest", "Encryption in transit", "VPC deployment", "Audit logging", "Compliance alignment"], default=["IAM-based access", "VPC deployment"])
st.divider()

# --- 8. DYNAMIC TIMELINE ---
st.header("📅 8. Timeline & Phasing")
poc_dur_val = st.selectbox("PoC Duration:", ["2 weeks", "4 weeks", "6 weeks", "8 weeks", "Custom"])
num_weeks = 0
if poc_dur_val == "Custom": num_weeks = st.number_input("Enter weeks:", 1, 12, 4)
else: num_weeks = int(poc_dur_val.split()[0])

base_cols = ["Phase", "Task"]
wk_cols = [f"Wk{i}" for i in range(1, num_weeks + 1)]
new_cols = base_cols + wk_cols

if list(st.session_state.timeline_phases.columns) != new_cols:
    old_df = st.session_state.timeline_phases
    new_df = old_df[["Phase", "Task"]].copy() if "Phase" in old_df.columns else pd.DataFrame(columns=["Phase", "Task"])
    for wk in wk_cols:
        if wk in old_df.columns: new_df[wk] = old_df[wk]
        else: new_df[wk] = ""
    st.session_state.timeline_phases = new_df

st.session_state.timeline_phases = st.data_editor(st.session_state.timeline_phases, num_rows="dynamic", key=f"ed_t_{num_weeks}")
st.divider()

st.header("💰 9. Costing Inputs & Ownership")
st.info(f"Calculator Link: {CALCULATOR_LINKS.get(sow_key, 'https://calculator.aws')}")
ownership = st.selectbox("Cost Ownership:", ["Funded by AWS", "Funded by Partner", "Funded by Customer", "Shared"], index=2)
st.divider()

st.header("🏁 10. Final Outputs")
delivs = st.multiselect("Deliverables:", ["PoC architecture", "Working demo", "SOW document", "Cost estimate", "Next-phase proposal"], default=["Working demo", "SOW document"])
nxt = st.multiselect("Next Steps:", ["Production proposal", "Scaling roadmap", "Security review", "Performance optimization", "Model fine-tuning"], default=["Production proposal", "Scaling roadmap"])

# --- GENERATION ---
if st.button("✨ Generate Full SOW", type="primary", use_container_width=True):
    with st.spinner("Applying structured standards..."):
        def get_md(df):
            headers = "| " + " | ".join(df.columns) + " |"
            sep = "| " + " | ".join(["---"] * len(df.columns)) + " |"
            rows = ["| " + " | ".join(str(val) for val in row) + " |" for _, row in df.iterrows()]
            return "\n".join([headers, sep] + rows)
            
        cost_info = SOW_COST_TABLE_MAP.get(sow_key, {})
        cost_table = "| System | Infra Cost | AWS Cost Calculator Link |\n| --- | --- | --- |\n"
        for k,v in cost_info.items(): 
            label = "POC" if k == "poc_cost" else "Production" if k == "prod_cost" else k
            cost_table += f"| {label} | {v} | Link |\n"
        
        prompt = f"""
        Generate a professional enterprise SOW for {sow_key}. 
        STRICT FORMAT: Syngene 5-Section standard. Main headings MUST be BOLD.

        # 2  Project Overview
        ## 2.1 Objective
        (Rewrite {biz_objective} formally).

        ## 2.2 Project Sponsor(s) / Stakeholder(s) / Project Team
        ### Partner Executive Sponsor
        {get_md(st.session_state.stakeholders["Partner"])}
        ### Customer Executive Sponsor
        {get_md(st.session_state.stakeholders["Customer"])}
        ### AWS Executive Sponsor
        {get_md(st.session_state.stakeholders["AWS"])}
        ### Project Escalation Contacts
        {get_md(st.session_state.stakeholders["Escalation"])}

        ## 2.3 Assumptions & Dependencies
        ### Dependencies: 
        The following items are required from the customer prior to the commencement of the POC:
        1. Access to {', '.join(sel_deps)}
        2. Data characteristics: {data_meta}

        ### Assumptions: 
        - {', '.join(sel_ass)} {custom_ass}

        ## 2.4 PoC Success Criteria
        Success outcomes (BOLD headers with sub-bullets):
        1. *Capability Validation*
           o Successful demonstration of {', '.join(sel_caps)}
        2. *Result Quality*
           o Target metrics for {', '.join(sel_dims)}
        3. *Validation Outcome*
           o {val_req}.

        # 3  Scope of Work - Technical Project Plan
        Scope of Work:
        A. Infrastructure Setup
        o Configure {', '.join(ai_svcs)}.
        o Setup minimal AWS backend for data ingestion.
        B. Core Component Development
        o Implement RAG logic and functional flows: {', '.join(sel_caps)}.
        o Integration with {', '.join(sel_ints)}.
        C. UI Development (Streamlit)
        o Build interactive interface for outcome visualization.
        D. Testing & Feedback
        o Stakeholder iterative reviews and formal sign-off.

        Development Timelines: 
        [TIMELINE_PLACEHOLDER]

        # 4  Solution Architecture / Architectural Diagram
        No description 
        Only Disclaimer- *Specifics to be discussed basis POC 
        
        {cost_table}

        # 5  RESOURCES & COST ESTIMATES
        The POC Development costs are funded jointly by AWS and Oneture.
        Cost ownership: {ownership}.
        Deliverables: {', '.join(delivs)}.
        """
        payload = {"contents": [{"parts": [{"text": prompt}]}], "systemInstruction": {"parts": [{"text": "AWS Solutions Architect. Adhere to sections 2-5 numbering. BOLD main headings. Use letters A, B, C for sub-headings in Section 3."}]}}
        res, err = call_gemini_with_retry(payload, api_key_input=api_key)
        if res:
            st.session_state.generated_sow = res.json()['candidates'][0]['content']['parts'][0]['text']
            st.rerun()
        else: st.error(err)

# --- REVIEW & EXPORT ---
if st.session_state.generated_sow:
    st.divider(); tab_e, tab_p = st.tabs(["✍️ Editor", "📄 Visual Preview"])
    with tab_e: st.session_state.generated_sow = st.text_area("Modify SOW:", st.session_state.generated_sow, height=600)
    with tab_p:
        st.markdown('<div class="sow-preview">', unsafe_allow_html=True)
        preview_toc = "## Table of Contents\n- 2  Project Overview\n- 3  Scope of Work - Technical Project Plan\n- 4  Solution Architecture / Architectural Diagram\n- 5  RESOURCES & COST ESTIMATES\n\n---\n"
        # Removing asterisks from generated content for cleaner look
        clean_content = st.session_state.generated_sow.replace("**", "")
        full_content = preview_toc + clean_content
        calc_url_p = CALCULATOR_LINKS.get(sow_key, "https://calculator.aws/")
        
        def render_html_timeline(df):
            html = '<table class="timeline-table"><thead><tr>'
            for col in df.columns: html += f'<th>{col}</th>'
            html += '</tr></thead><tbody>'
            for _, row in df.iterrows():
                html += '<tr>'
                for col in df.columns:
                    val = str(row[col])
                    shade = 'shaded-cell' if val.strip().upper() == "X" else ''
                    disp = "" if val.strip().upper() == "X" else val
                    html += f'<td class="{shade}">{disp}</td>'
                html += '</tr>'
            html += '</tbody></table>'
            return html

        s3_pattern = r"#\s*3\s+Scope of Work"
        s4_pattern = r"#\s*4\s+Solution Architecture"
        
        parts = re.split(f"({s3_pattern}|{s4_pattern})", full_content, flags=re.IGNORECASE)
        
        for p in parts:
            if re.match(s3_pattern, p, re.IGNORECASE): 
                st.markdown(f"*{p}*", unsafe_allow_html=True)
            elif re.match(s4_pattern, p, re.IGNORECASE): 
                st.markdown(f"*{p}*", unsafe_allow_html=True)
                diag_out = SOW_DIAGRAM_MAP.get(sow_key.strip())
                if diag_out and os.path.exists(diag_out): 
                    st.image(diag_out, caption=f"{sow_key} Architecture Diagram")
            else:
                final_p = p.replace("Link", f'<a href="{calc_url_p}" target="_blank">Link</a>')
                if "Development Timelines:" in final_p:
                    sub_parts = final_p.split("Development Timelines:")
                    st.markdown(sub_parts[0], unsafe_allow_html=True)
                    st.markdown("*Development Timelines:*", unsafe_allow_html=True)
                    st.write(render_html_timeline(st.session_state.timeline_phases), unsafe_allow_html=True)
                    if len(sub_parts) > 1: st.markdown(sub_parts[1], unsafe_allow_html=True)
                else:
                    st.markdown(final_p, unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    if st.button("💾 Prepare Microsoft Word"):
        branding = {"sow_name": sow_key, "customer_logo_bytes": customer_logo.getvalue() if customer_logo else None, "doc_date_str": doc_date.strftime("%d %B %Y")}
        docx_data = create_docx_logic(st.session_state.generated_sow, branding, sow_key, st.session_state.timeline_phases)
        st.download_button("📥 Download SOW (.docx)", docx_data, f"SOW_{sow_key.replace(' ', '_')}.docx", use_container_width=True)
